# SPDX-License-Identifier: MIT
from __future__ import annotations

import argparse
import zipfile
from pathlib import Path

from docx import Document

from article_schema import body_parts, load_article, nonspace_count, validate_privacy, validate_target_count


def validate_docx(data: dict, path: Path) -> dict[str, int]:
    if not path.is_file() or path.stat().st_size == 0:
        raise ValueError(f"missing or empty DOCX: {path}")
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        if bad:
            raise ValueError(f"corrupt DOCX member: {bad}")

    document = Document(path)
    if not document.paragraphs or document.paragraphs[0].text != data["title"]:
        raise ValueError("DOCX title does not match JSON")
    actual_parts = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.style.name in {"Article Section", "Article Body"} and paragraph.text.strip()
    ]
    expected_parts = body_parts(data)
    if actual_parts != expected_parts:
        raise ValueError("DOCX body does not match JSON")

    properties = document.core_properties
    if properties.author or properties.last_modified_by or properties.comments:
        raise ValueError("DOCX contains author, last-modified-by, or comments metadata")

    links = sum(1 for relation in document.part.rels.values() if "hyperlink" in relation.reltype)
    if links != len(data.get("sources", [])):
        raise ValueError(f"expected {len(data.get('sources', []))} hyperlinks, found {links}")
    if document.tables or document.inline_shapes:
        raise ValueError("default DOCX must not contain tables or inline images")
    return {"body_chars": nonspace_count(actual_parts), "paragraphs": len(document.paragraphs), "hyperlinks": links}


def main() -> int:
    parser = argparse.ArgumentParser(description="Validate article JSON and optional rendered DOCX.")
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--docx", type=Path)
    args = parser.parse_args()
    data = load_article(args.input)
    count = validate_target_count(data)
    validate_privacy(data)
    report = {"body_chars": count}
    if args.docx:
        report.update(validate_docx(data, args.docx))
    print("PASS " + " ".join(f"{key}={value}" for key, value in report.items()))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
