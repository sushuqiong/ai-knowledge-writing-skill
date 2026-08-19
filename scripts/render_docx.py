# SPDX-License-Identifier: MIT
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Cm, Pt, RGBColor

from article_schema import load_article, validate_privacy, validate_target_count


FONT = "Microsoft YaHei"


def set_font(run, size: float, bold: bool = False, color: tuple[int, int, int] = (31, 41, 55)) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)


def ensure_style(document: Document, name: str, size: float, bold: bool = False):
    styles = document.styles
    style = styles[name] if name in styles else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = FONT
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    return style


def add_hyperlink(paragraph, label: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0F766E")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    run.append(properties)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def render(data: dict, output: Path) -> None:
    validate_target_count(data)
    validate_privacy(data)
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    ensure_style(document, "Article Section", 13, True)
    ensure_style(document, "Article Body", 10.5)
    ensure_style(document, "Article Disclaimer", 9)
    ensure_style(document, "Article Sources", 11, True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_font(title.add_run(data["title"]), 18, True, (24, 52, 48))

    if data.get("subtitle"):
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(16)
        set_font(subtitle.add_run(data["subtitle"]), 9, False, (88, 99, 96))

    for section_data in data["sections"]:
        if section_data.get("heading", "").strip():
            heading = document.add_paragraph(style="Article Section")
            heading.paragraph_format.keep_with_next = True
            heading.paragraph_format.space_before = Pt(10)
            heading.paragraph_format.space_after = Pt(5)
            set_font(heading.add_run(section_data["heading"].strip()), 13, True, (15, 118, 110))
        for paragraph_text in section_data["paragraphs"]:
            paragraph = document.add_paragraph(style="Article Body")
            paragraph.paragraph_format.first_line_indent = Cm(0.74)
            paragraph.paragraph_format.line_spacing = 1.45
            paragraph.paragraph_format.space_after = Pt(7)
            set_font(paragraph.add_run(paragraph_text.strip()), 10.5)

    if data.get("disclaimer"):
        paragraph = document.add_paragraph(style="Article Disclaimer")
        paragraph.paragraph_format.space_before = Pt(10)
        set_font(paragraph.add_run(data["disclaimer"].strip()), 9, False, (110, 84, 45))

    if data.get("sources"):
        heading = document.add_paragraph(style="Article Sources")
        heading.paragraph_format.space_before = Pt(14)
        set_font(heading.add_run("资料来源"), 11, True, (24, 52, 48))
        for index, source in enumerate(data["sources"], start=1):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(3)
            prefix = paragraph.add_run(f"{index}. {source['label']}：")
            set_font(prefix, 8.5)
            add_hyperlink(paragraph, source["url"], source["url"])
            if source.get("accessed"):
                suffix = paragraph.add_run(f"（访问：{source['accessed']}）")
                set_font(suffix, 8.5, False, (88, 99, 96))

    properties = document.core_properties
    properties.title = data["title"]
    properties.subject = "Public knowledge explainer"
    properties.author = ""
    properties.last_modified_by = ""
    properties.comments = ""
    properties.keywords = ""
    properties.category = ""

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description="Render a validated public article as DOCX.")
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    data = load_article(args.input)
    render(data, args.output)
    print(f"WROTE {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
