# SPDX-License-Identifier: MIT
from __future__ import annotations

import copy
import json
import shutil
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
sys.path.insert(0, str(SCRIPTS))

from article_schema import (  # noqa: E402
    ArticleValidationError,
    body_parts,
    load_article,
    nonspace_count,
    validate_article_data,
    validate_privacy,
    validate_target_count,
)
from render_docx import render  # noqa: E402
from validate_article import validate_docx  # noqa: E402
from validate_public_package import (  # noqa: E402
    PackageValidationError,
    validate_evals,
    validate_repository,
    validate_skill,
)


def article_data() -> dict:
    return json.loads((ROOT / "templates/article.example.json").read_text(encoding="utf-8"))


class PublicPackageTests(unittest.TestCase):
    def copy_repository(self, destination: Path) -> Path:
        target = destination / "repo"
        shutil.copytree(ROOT, target, ignore=shutil.ignore_patterns(".git", "dist", "__pycache__", "*.pyc"))
        return target

    def test_01_repository_is_valid(self):
        result = validate_repository(ROOT)
        self.assertEqual(result["eval_cases"], 24)

    def test_02_skill_frontmatter(self):
        validate_skill(ROOT)

    def test_03_broken_markdown_link_is_rejected(self):
        with tempfile.TemporaryDirectory() as directory:
            root = self.copy_repository(Path(directory))
            with (root / "README.md").open("a", encoding="utf-8") as handle:
                handle.write("\n[missing](missing-file.md)\n")
            with self.assertRaises(PackageValidationError):
                validate_repository(root)

    def test_04_private_path_is_rejected(self):
        with tempfile.TemporaryDirectory() as directory:
            root = self.copy_repository(Path(directory))
            marker = "C:" + "\\" + "Users" + "\\" + "private-name" + "\\" + "note.txt"
            with (root / "README.md").open("a", encoding="utf-8") as handle:
                handle.write("\n" + marker + "\n")
            with self.assertRaises(PackageValidationError):
                validate_repository(root)

    def test_05_personal_email_is_rejected(self):
        with tempfile.TemporaryDirectory() as directory:
            root = self.copy_repository(Path(directory))
            marker = "person" + "@" + "example.com"
            with (root / "README.md").open("a", encoding="utf-8") as handle:
                handle.write("\n" + marker + "\n")
            with self.assertRaises(PackageValidationError):
                validate_repository(root)

    def test_06_credential_assignment_is_rejected(self):
        with tempfile.TemporaryDirectory() as directory:
            root = self.copy_repository(Path(directory))
            marker = "api_" + "key=" + "abcdefghijklmnop"
            with (root / "README.md").open("a", encoding="utf-8") as handle:
                handle.write("\n" + marker + "\n")
            with self.assertRaises(PackageValidationError):
                validate_repository(root)

    def test_07_eval_suite_has_24_cases(self):
        self.assertEqual(validate_evals(ROOT), 24)


class ArticleToolchainTests(unittest.TestCase):
    def test_08_valid_article_schema(self):
        data = article_data()
        validate_article_data(data)
        self.assertEqual(load_article(ROOT / "templates/article.example.json")["schema_version"], 1)

    def test_09_body_count_is_in_range(self):
        data = article_data()
        count = validate_target_count(data)
        self.assertEqual(count, nonspace_count(body_parts(data)))
        self.assertGreaterEqual(count, data["target_chars"]["min"])
        self.assertLessEqual(count, data["target_chars"]["max"])

    def test_10_invalid_source_scheme_is_rejected(self):
        data = copy.deepcopy(article_data())
        data["sources"][0]["url"] = "file:///private/source"
        with self.assertRaises(ArticleValidationError):
            validate_article_data(data)

    def test_11_article_privacy_finding_is_rejected(self):
        data = copy.deepcopy(article_data())
        data["title"] = "C:" + "\\" + "Users" + "\\" + "private-name"
        with self.assertRaises(ArticleValidationError):
            validate_privacy(data)

    def test_12_rendered_docx_round_trip(self):
        data = article_data()
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "article.docx"
            render(data, output)
            report = validate_docx(data, output)
            self.assertEqual(report["body_chars"], validate_target_count(data))

    def test_13_docx_metadata_is_cleared(self):
        data = article_data()
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "article.docx"
            render(data, output)
            properties = Document(output).core_properties
            self.assertEqual(properties.author, "")
            self.assertEqual(properties.last_modified_by, "")
            self.assertEqual(properties.comments, "")

    def test_14_unicode_and_hyperlinks_are_preserved(self):
        data = article_data()
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "article.docx"
            render(data, output)
            document = Document(output)
            visible = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("材料只是起点", visible)
            links = sum(1 for relation in document.part.rels.values() if "hyperlink" in relation.reltype)
            self.assertEqual(links, 1)


if __name__ == "__main__":
    unittest.main()
